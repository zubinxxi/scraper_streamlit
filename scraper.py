import streamlit as st
import requests as rq
import re
from bs4 import BeautifulSoup
from docx import Document
from io import BytesIO


def scraper():
    st.subheader("Wikipedia Web Scraper")

    search_term = st.text_input("", placeholder="Buscar en Wikipedia", autocomplete="off")

    if search_term:
        base_url = "https://es.wikipedia.org/wiki/"

        # Define un User-Agent común de navegador web
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }

        url = base_url + search_term.replace(" ", "_")
        st.write(f"URL generada: {url}")

        try:
            response = rq.get(url, headers=headers)
            response.raise_for_status()
            st.write("**Respuesta del servidor Success (200)**")
            soup = BeautifulSoup(response.text, 'html.parser')
            title = soup.find('h1').text
            st.write(f"**Título del articulo: {title}**")
            #paragraph = soup.find('p').text
            #st.write(paragraph)

            word_text = []

            paragraphs = soup.find_all('p')
            for paragraph in paragraphs:
                text = re.sub(r'\[\d+\]', '', paragraph.text)
                word_text.append(text)
                st.write(text)

            if st.button("Generar Archivo de Word", type="primary"):
                doc = Document()
                doc.add_heading(title, level=1)
                for p in word_text:
                    doc.add_paragraph(p)
                
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.download_button(
                    "Descargar documento",
                    icon=":material/download:",
                    data=buffer,
                    file_name="documento.docx",
                    mime="application/vnd.openxmlformats-officedocumet.wordprocessingml.document"
                )

        except rq.exceptions.RequestException as e:
            st.error(f"Error al realizar solicitud: {e}")
        except Exception as e:
            print(e)
            st.error(f"Error al analizar el html: {e}")    

